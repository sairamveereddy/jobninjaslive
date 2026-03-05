"""
Resume Parser — PDF/DOCX -> text -> skills -> inferred title.
"""
import io, json, re

try:
    import fitz
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

SKILLS_DB = {
    "languages":    ["python","javascript","typescript","java","go","golang","rust","c++","c#","ruby","swift","kotlin","php","scala","r","sql","bash","dart","elixir"],
    "frontend":     ["react","vue","angular","svelte","nextjs","nuxtjs","gatsby","html","css","tailwind","bootstrap","sass","webpack","vite","redux","graphql","jquery"],
    "backend":      ["node","nodejs","express","fastapi","django","flask","spring","rails","laravel","nestjs","fastify","gin","fiber","actix","grpc","microservices"],
    "databases":    ["postgresql","mysql","mongodb","redis","elasticsearch","dynamodb","cassandra","sqlite","firebase","supabase","pinecone","neo4j","snowflake","bigquery"],
    "cloud_devops": ["aws","gcp","azure","docker","kubernetes","k8s","terraform","ansible","jenkins","github actions","circleci","helm","prometheus","grafana","datadog","nginx","linux"],
    "ml_ai":        ["pytorch","tensorflow","keras","scikit-learn","pandas","numpy","huggingface","langchain","openai","machine learning","deep learning","nlp","computer vision","data science","mlops","spark","airflow","dbt","xgboost","bert","rag"],
    "tools":        ["git","github","gitlab","jira","confluence","figma","postman","tableau","powerbi","looker"],
}

ALL_SKILLS = [s for skills in SKILLS_DB.values() for s in skills]

SKILL_WEIGHTS = {
    "languages": 15, "frontend": 12, "backend": 12,
    "databases": 8, "cloud_devops": 10, "ml_ai": 14, "tools": 5,
}

def extract_text_pdf(data: bytes) -> str:
    if not HAS_PDF:
        raise RuntimeError("Run: pip install pymupdf")
    doc = fitz.open(stream=data, filetype="pdf")
    return "\n".join(p.get_text("text") for p in doc).strip()

def extract_text_docx(data: bytes) -> str:
    if not HAS_DOCX:
        raise RuntimeError("Run: pip install python-docx")
    doc = DocxDocument(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)

def extract_skills(text: str) -> list:
    tl = text.lower()
    return list(dict.fromkeys(
        s for s in ALL_SKILLS
        if re.search(r"\b" + re.escape(s) + r"\b", tl)
    ))

def infer_title(text: str, skills: list) -> str:
    tl = text.lower()
    patterns = [
        r"(staff|principal|senior|lead|junior)?\s*(software|backend|frontend|full.?stack|ml|ai|data|devops|platform|mobile|ios|android|security|cloud|sre)\s*(engineer|developer|architect|scientist|analyst)",
        r"(engineering|product|technical)\s*(manager|director|lead)",
        r"(data|ml|ai|analytics)\s*(scientist|engineer|analyst)",
        r"(ux|ui|product)\s*(designer|researcher)",
    ]
    for pat in patterns:
        m = re.search(pat, tl)
        if m:
            return re.sub(r"\s+", " ", m.group(0).title().strip())
    s = set(skills)
    if any(x in s for x in ["pytorch","tensorflow","machine learning","deep learning"]):
        return "Machine Learning Engineer"
    if any(x in s for x in ["react","vue","angular","nextjs"]):
        return "Frontend Engineer"
    if any(x in s for x in ["fastapi","django","flask","express","spring"]):
        return "Backend Engineer"
    if any(x in s for x in ["kubernetes","docker","terraform"]):
        return "DevOps Engineer"
    if any(x in s for x in ["pandas","spark","dbt","airflow"]):
        return "Data Engineer"
    return "Software Engineer"

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        text = extract_text_pdf(file_bytes)
    elif ext in ("docx","doc"):
        text = extract_text_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
    if len(text.strip()) < 50:
        raise ValueError("Could not extract text from resume")
    skills = extract_skills(text)
    return {
        "text": text, "skills": skills,
        "title": infer_title(text, skills),
        "word_count": len(text.split()),
        "skill_count": len(skills),
    }
